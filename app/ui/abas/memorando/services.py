# -*- coding: utf-8 -*-
"""Services para operações de memorandos - Versão SQL Server"""

import io
import os
import subprocess
import sys
import tempfile
from datetime import datetime

from docx import Document
import pyodbc

from services.database import get_connection


def _to_bytes(valor) -> bytes:
    """Converte dados binários para bytes de forma segura."""
    if valor is None:
        return b""
    if isinstance(valor, memoryview):
        return bytes(valor)
    if isinstance(valor, bytearray):
        return bytes(valor)
    return bytes(valor) if isinstance(valor, (bytes, bytearray)) else b""


def _fmt_data(valor) -> str:
    """Formata date/datetime para DD/MM/YYYY."""
    if not valor:
        return ""
    if hasattr(valor, "strftime"):
        return valor.strftime("%d/%m/%Y")
    return str(valor)


def _normalizar_data(data_str: str) -> str:
    """Normaliza data para DD/MM/AAAA"""
    data_str = data_str.strip()
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(data_str, fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return data_str


def _substituir_runs(paragraph, substituicoes: dict) -> None:
    """Substitui placeholders preservando a formatação"""
    texto_original = "".join(run.text for run in paragraph.runs)
    texto_novo = texto_original

    for placeholder, valor in substituicoes.items():
        texto_novo = texto_novo.replace(placeholder, str(valor))

    if texto_novo == texto_original:
        return

    runs = paragraph.runs
    if not runs:
        return

    runs[0].text = texto_novo
    for run in runs[1:]:
        run.text = ""


class MemorandoService:

    @staticmethod
    def buscar_modelo(modelo_id: int = 11) -> bytes:
        """Busca o template .docx da tabela `documento` (SQL Server)"""
        conn = get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT conteudo FROM documento WHERE id = ?", (modelo_id,))
            row = cursor.fetchone()
            cursor.close()

            if not row or row[0] is None:
                raise ValueError(
                    f"Modelo id={modelo_id} não encontrado na tabela documento.\n"
                    "Execute o script insert_modelo_memorando.sql primeiro."
                )
            return _to_bytes(row[0])
        finally:
            conn.close()

    @staticmethod
    def criar_memorando_word(numero: str, data: str, unloc: str,
                             memo_entrada: str = "",
                             nomes: str = "—",
                             qtda: str = "0") -> bytes:
        """Cria documento Word com placeholders substituídos"""
        data_formatada = _normalizar_data(data)

        modelo_bytes = MemorandoService.buscar_modelo()
        doc = Document(io.BytesIO(modelo_bytes))

        substituicoes = {
            "(num)": numero,
            "(data)": data_formatada,
            "(muni)": unloc.upper(),
            "(memos)": memo_entrada,
            "(nomes)": nomes or "—",
            "(qtda)": qtda or "0",
        }

        for paragraph in doc.paragraphs:
            _substituir_runs(paragraph, substituicoes)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _substituir_runs(paragraph, substituicoes)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def salvar_memorando_bd(dados: dict) -> dict:
        """Salva memorando no banco SQL Server"""
        conn = get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO memorandos (
                    numero, descricao, data_emissao,
                    unloc, municipio, memo_entrada,
                    quantidade, word_conteudo, usuario
                )
                OUTPUT INSERTED.id, INSERTED.numero, INSERTED.unloc,
                       INSERTED.data_emissao, INSERTED.criado_em, INSERTED.usuario
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                dados["numero"],
                f"Memorando {dados['numero']}/{dados['unloc']}",
                dados["data_emissao"],
                dados["unloc"],
                dados["unloc"],
                dados.get("memo_entrada", ""),
                dados.get("quantidade", 0),
                pyodbc.Binary(dados["word_conteudo"]),
                dados.get("usuario", "sistema"),
            ))

            row = cursor.fetchone()
            cursor.close()
            conn.commit()

            return {
                "id": row[0],
                "numero": row[1],
                "unloc": row[2],
                "data_emissao": row[3],
                "criado_em": row[4],
                "usuario": row[5],
            }
        except Exception:
            conn.rollback()
            raise
        finally:
            conn.close()

    @staticmethod
    def buscar_memorandos(termo_pesquisa: str = "", limite: int = 50) -> list:
        """Busca histórico de memorandos com termo simples"""
        conn = get_connection()
        try:
            cursor = conn.cursor()

            if termo_pesquisa.strip():
                cursor.execute("""
                    SELECT TOP (?)
                           id, numero, unloc, data_emissao,
                           criado_em, usuario, descricao
                    FROM memorandos
                    WHERE numero LIKE ? OR unloc LIKE ?
                    ORDER BY criado_em DESC
                """, (limite, f"%{termo_pesquisa}%", f"%{termo_pesquisa}%"))
            else:
                cursor.execute("""
                    SELECT TOP (?)
                           id, numero, unloc, data_emissao,
                           criado_em, usuario, descricao
                    FROM memorandos
                    ORDER BY criado_em DESC
                """, (limite,))

            rows = cursor.fetchall()
            cursor.close()

            return [
                {
                    "id": row[0],
                    "numero": row[1],
                    "unloc": row[2],
                    "data_emissao": _fmt_data(row[3]),
                    "criado_em": _fmt_data(row[4]),
                    "usuario": row[5],
                    "descricao": row[6],
                }
                for row in rows
            ]
        finally:
            conn.close()

    @staticmethod
    def buscar_memorandos_com_filtros(
        termo: str = "",
        codigo_municipio: str = None,
        ano: str = None,
        ordem: str = "Recentes",
        limite: int = 200
    ) -> list:
        """
        Busca memorandos com filtros avançados.
        """
        conn = get_connection()
        try:
            cursor = conn.cursor()

            conditions = []
            params = []

            if termo and termo.strip():
                conditions.append("(numero LIKE ? OR unloc LIKE ?)")
                params.extend([f"%{termo.strip()}%", f"%{termo.strip()}%"])

            if codigo_municipio and codigo_municipio.strip():
                conditions.append("unloc LIKE ?")
                params.append(f"{codigo_municipio.strip()}%")

            if ano and str(ano).strip():
                conditions.append("YEAR(data_emissao) = ?")
                params.append(int(str(ano).strip()))

            where_clause = ""
            if conditions:
                where_clause = "WHERE " + " AND ".join(conditions)

            order_direction = "DESC" if ordem == "Recentes" else "ASC"

            sql = f"""
                SELECT TOP ({limite})
                       id, numero, unloc, data_emissao,
                       criado_em, usuario, descricao
                FROM memorandos
                {where_clause}
                ORDER BY criado_em {order_direction}
            """

            cursor.execute(sql, params)
            rows = cursor.fetchall()
            cursor.close()

            return [
                {
                    "id": row[0],
                    "numero": row[1],
                    "unloc": row[2],
                    "data_emissao": _fmt_data(row[3]),
                    "criado_em": _fmt_data(row[4]),
                    "usuario": row[5],
                    "descricao": row[6],
                }
                for row in rows
            ]
        finally:
            conn.close()

    @staticmethod
    def listar_anos_disponiveis() -> list:
        """Retorna lista de anos distintos presentes nos memorandos (ordem decrescente)."""
        conn = get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT DISTINCT YEAR(data_emissao) AS ano
                FROM memorandos
                WHERE data_emissao IS NOT NULL
                ORDER BY ano DESC
            """)
            rows = cursor.fetchall()
            cursor.close()
            return [str(row[0]) for row in rows if row[0] is not None]
        finally:
            conn.close()

    @staticmethod
    def buscar_memorando_por_id(memorando_id: int):
        """Busca memorando por ID"""
        conn = get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, numero, unloc, data_emissao,
                       word_conteudo, usuario, descricao
                FROM memorandos
                WHERE id = ?
            """, (memorando_id,))

            row = cursor.fetchone()
            cursor.close()

            if not row:
                return None

            return {
                "id": row[0],
                "numero": row[1],
                "unloc": row[2],
                "data_emissao": row[3],
                "word_conteudo": _to_bytes(row[4]),
                "usuario": row[5],
                "descricao": row[6],
            }
        finally:
            conn.close()

    @staticmethod
    def abrir_arquivo_temp(conteudo_word: bytes, sufixo: str = ".docx") -> str:
        """Abre arquivo temporário com o documento"""
        conteudo = _to_bytes(conteudo_word)

        with tempfile.NamedTemporaryFile(delete=False, suffix=sufixo) as tmp:
            tmp.write(conteudo)
            caminho = tmp.name

        if sys.platform.startswith("win"):
            os.startfile(caminho)
        elif sys.platform.startswith("darwin"):
            subprocess.call(["open", caminho])
        else:
            subprocess.call(["xdg-open", caminho])

        return caminho