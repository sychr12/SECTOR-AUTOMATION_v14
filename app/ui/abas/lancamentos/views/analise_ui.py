# -*- coding: utf-8 -*-
"""Interface de Análise de Processos — design refinado + bugs corrigidos."""
import smtplib
import threading
from datetime import datetime
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from PyQt6.QtWidgets import *

from PyQt6.QtWidgets import *
from PyQt6.QtCore import Qt
from PyQt6.QtGui import *

from app.theme import AppTheme
from ui.base_ui import BaseUI
from ..backend.controller import AnaliseController
from ..backend.services import AnaliseService
from .devolucao_view import DevolucaoPopup
from .historico import HistoricoView

# ── Paleta ───────────────────────────────────────────────────────────────
_VERDE    = "#22c55e"
_VERDE_H  = "#16a34a"
_AZUL     = "#3b82f6"
_AZUL_H   = "#2563eb"
_VERMELHO = "#ef4444"
_VERM_H   = "#dc2626"
_MUTED    = "#64748b"
_AMBER    = "#f59e0b"


def _badge(parent, texto: str, cor: str) -> ctk.CTkLabel:
    """Mini badge colorido."""
    frame = ctk.CTkFrame(parent, fg_color=cor, corner_radius=6)
    lbl = ctk.CTkLabel(frame, text=texto,
                       font=("Segoe UI", 10, "bold"),
                       text_color="#ffffff", padx=8, pady=2)
    lbl.pack()
    return frame


class AnaliseUI(BaseUI):

    def __init__(self, master, usuario: str, on_refresh_lancamento=None,
                 on_refresh_devolucao=None):
        super().__init__(master)
        self.usuario    = usuario
        self.service    = AnaliseService()
        self.controller = AnaliseController(
            usuario, self.service.repo,
            self.service.repo_ap,
            self.service.sefaz_repo
        )
        self._on_refresh_lancamento = on_refresh_lancamento
        self._on_refresh_devolucao  = on_refresh_devolucao
        self.var_urgencia   = ctk.BooleanVar(value=False)
        self.registros_vars: dict[int, ctk.BooleanVar] = {}
        self._todos_registros: list = []   # cache para filtragem local

        self._motivos = self._carregar_motivos()
        self.configure(fg_color=AppTheme.BG_APP)

        # ── Notebook ────────────────────────────────────────────────
        self.nb = ctk.CTkTabview(
            self,
            fg_color=AppTheme.BG_CARD,
            segmented_button_fg_color=AppTheme.BG_INPUT,
            segmented_button_selected_color=_VERDE,
            segmented_button_selected_hover_color=_VERDE_H,
            segmented_button_unselected_color=AppTheme.BG_INPUT,
            segmented_button_unselected_hover_color=AppTheme.BG_CARD,
            text_color=AppTheme.TXT_MAIN,
            text_color_disabled=_MUTED,
        )
        self.nb.pack(fill="both", expand=True, padx=0, pady=0)

        self.tab_analise   = self.nb.add("  Análise de Processos  ")
        self.tab_adicionar = self.nb.add("  Adicionar Registro  ")

        self._build_analise()
        self._build_adicionar()

        # Carregar dados em thread para não travar UI
        threading.Thread(target=self._carregar_thread, daemon=True).start()

        if hasattr(master, "protocol"):
            master.protocol("WM_DELETE_WINDOW", self._on_close)

    # ────────────────────────────────────────────────────────────────
    # Motivos de devolução
    # ────────────────────────────────────────────────────────────────
    def _carregar_motivos(self) -> dict:
        """Motivos de devolução embutidos no código."""
        return {
            "Endereço": [
                "NOME DA COMUNIDADE ESTÁ DIVERGENTE ENTRE FAC E DECLARAÇÃO;",
                "CORRIGIR A COORDENADAS, NÃO SÃO DO ENDEREÇO DESCRITO;",
                "ENDEREÇO INCOMPLETO, FALTA A COMUNIDADE OU SE É PRODUTOR ISOLADO;",
                "ENVIAR COMO ALTERAÇÃO DE ENDEREÇO;",
                "FALTA O KM DA RODOVIA;",
                "FALTA O KM DA ESTRADA;",
                "FALTA O KM DO RAMAL;",
                "FALTA O KM DA VICINAL;",
                "COORDENADAS E ENDEREÇO ESTÃO DIVERGENTES;",
                "ENDEREÇO ESTÁ DIVERGENTE ENTRE FAC E DECLARAÇÃO;",
                "COORDENADAS INVÁLIDAS, CORRIGIR;",
                "O ENDEREÇO DA ÁREA DA PESCA ESTÁ DIVERGENTE ENTRE FAC E DECLARAÇÃO;",
                "COORDENADAS E ENDEREÇO ESTÃO DIVERGENTES, CORRIGIR;",
                "FALTA A COMUNIDADE NA DECLARAÇÃO;",
                "ENDEREÇO DIVERGENTE DA CPP, CORRIGIR OU ENVIAR COMO ALTERAÇÃO;",
            ],
            "Documentos": [
                "ENVIAR OUTRO DOCUMENTO DE IDENTIFICAÇÃO, RG, OU HABILITAÇÃO, ETC;",
                "FALTA UM DOCUMENTO DE IDENTIFICAÇÃO COM FOTO;",
                "NA RG ELE NÃO ASSINA E ASSINOU NA FAC, ENTÃO APRESENTAR UM DOCUMENTO DE IDENTIDADE COM ASSINATURA;",
                "NÚMERO DO RG DIVERGENTE DO DOCUMENTO;",
                "NOME ESTÁ DIVERGENTE ENTRE RG E CPF, ATUALIZAR DOCUMENTOS;",
                "FRENTE DO RG ESTÁ ILEGÍVEL;",
                "FALTA O RG NA DECLARAÇÃO;",
                "COPIA DO RG ESTÁ ILEGÍVEL;",
                "ASSINATURA NA RG ESTÁ ILEGÍVEL;",
                "FALTA A AUTENTICAÇÃO DO CONTRATO;",
                "CORRIGIR CPF, DIVERGENTE DOS DOCUMENTOS PESSOAIS;",
                "CONTRATO DE COMODATO NÃO PODE SER POR TEMPO INDETERMINADO, PRECISA TER UM PRAZO ESTIPULADO;",
                "COM A NOVA DATA DE VALIDADE DA CPP DE 4 ANOS, OS CONTRATOS DEVEM TER O PRAZO DE NO MÍNIMO 4 ANOS DE VALIDADE;",
                "FALTA O TERMO DE EXTRATIVISMO;",
                "ENVIAR O COMPROVANTE DA SITUAÇÃO CADASTRAL NO CPF (DO SITE DA RECEITA FEDERAL);",
                "A SITUAÇÃO CADASTRAL ESTÁ SUSPENSA, REGULARIZAR O CPF PERANTE A RECEITA FEDERAL;",
                "ENVIAR CÓPIA DE UMA DOCUMENTAÇÃO COM FOTO E ASSINATURA;",
                "NECESSITA DO BOLETIM DE OCORRÊNCIA EMITIDO PELA POLICIA CIVIL;",
                "FALTA A FOLHA DE SOLICITAÇÃO DE BAIXA DE INSCRIÇÃO ESTADUAL;",
                "CARTEIRA DE TRABALHO TRANSCRITA A MÃO, NÃO É VÁLIDA;",
            ],
            "Cadastro": [
                "CORRIGIR O NOME DA PRODUTORA NA DECLARAÇÃO;",
                "JÁ POSSUI CADASTRO, ENVIAR COMO ALTERAÇÃO OU RENOVAÇÃO;",
                "CORRIGIR A DATA NA DECLARAÇÃO;",
                "ENVIAR TAMBÉM A FAC DIGITALIZADA;",
                "ENVIAR COMO ALTERAÇÃO, ENVIAR A CÓPIA DA CPP OU O BO;",
                "ATUALIZAR A FAC E A DECLARAÇÃO, ESTÃO COM DATA SUPERIOR A 6 MESES;",
                "NÚMERO DE CONTROLE DIVERGENTE DA CPP;",
                "NÚMERO DE CONTROLE DIVERGENTE NA DECLARAÇÃO;",
                "NÚMERO DE CONTROLE ESTÁ DIVERGENTE ENTRE DECLARAÇÃO E FAC;",
                "ESTÁ SEM O NÚMERO DE CONTROLE;",
                "NOME DA PROPRIEDADE ESTÁ DIVERGENTE ENTRE FAC E DECLARAÇÃO;",
                "NOME DA PROPRIEDADE ESTÁ DIVERGENTE DA CPP;",
                "ESPECIFICAR A CULTURA SECUNDÁRIA;",
                "AS DUAS ATIVIDADES ESTÃO DIVERGENTES ENTRE FAC E DECLARAÇÃO;",
                "COLOCAR A COMUNIDADE NA DECLARAÇÃO;",
                "COMUNIDADE ESTÁ DIVERGENTE ENTRE FAC E DECLARAÇÃO;",
                "ENVIAR COMO ALTERAÇÃO, HOUVE MUDANÇA NO SOBRENOME;",
                "ENVIAR COMO ALTERAÇÃO DE ATIVIDADE;",
                "FALTA O NOME DA PROPRIEDADE NA DECLARAÇÃO;",
                "FALTA O CARIMBO DO GERENTE;",
                "FALTA A ASSINATURA DO GERENTE;",
                "FALTA A ASSINATURA DO PRODUTOR NA FAC;",
                "NOME INCOMPLETO NA DECLARAÇÃO;",
                "NOME DO PRODUTOR ESTÁ DIVERGENTE ENTRE DECLARAÇÃO E FAC;",
                "CORRIGIR A SIGLA DO MUNICIPIO;",
                "ASSINATURA DO PRODUTOR ESTÁ DIVERGENTE DA RG;",
                "ATUALIZAR A DATA NA DECLARAÇÃO, ESTÁ SUPERIOR A 6 MESES;",
                "SEM ASSINATURA E CARIMBO DO GERENTE E TÉCNICO;",
                "APAGAR O NUMERO DA INSCRIÇÃO ESTADUAL NA FAC, É UMA INSCRIÇÃO;",
            ],
            "Pesca": [
                "REGISTRO DE PESCADOR NÃO ENCONTRADO NO SITE DO MINISTÉRIO DA PESCA E AGRICULTURA, APRESENTAR A CARTEIRA DE PESCADOR ATUAL;",
                "ESPECIFICAR AS ESPÉCIES E AS QUANTIDADES DE PEIXES E A ÁREA DA PISCICULTURA;",
                "FALTA A ÁREA DA PISCICULTURA;",
            ],
            "Simples Nacional": [
                "FALTA DAR A BAIXA;",
                "SIMPLES NACIONAL ATIVO;",
                "SIMPLES NACIONAL SUSPENSO;",
            ],
            "Animais": [
                "FALTA A QUANTIDADE DE FRANGOS;",
                "FALTA A QUANTIDADE E AS ESPÉCIES DE PEIXES;",
                "ENVIAR O COMPROVANTE DE SITUAÇÃO CADASTRAL NO CPF (DO SITE DA RECEITA FEDERAL);",
            ],
            "Outros": [],
        }

    def _abrir_motivos(self, categoria: str):
        """Abre popup com os motivos da categoria selecionada."""
        motivos = self._motivos.get(categoria, [])
        if not motivos:
            return

        popup = ctk.CTkToplevel(self)
        popup.title(f"Motivos — {categoria}")
        popup.geometry("520x420")
        popup.resizable(False, False)
        popup.configure(fg_color=AppTheme.BG_APP)
        popup.grab_set()
        popup.after(0, lambda: popup.geometry(
            f"520x420+"
            f"{popup.winfo_screenwidth()  // 2 - 260}+"
            f"{popup.winfo_screenheight() // 2 - 210}"
        ))

        ctk.CTkLabel(popup, text=f"Selecione o motivo — {categoria}",
                     font=("Segoe UI", 13, "bold"),
                     text_color=AppTheme.TXT_MAIN
                     ).pack(padx=24, pady=(18, 10), anchor="w")

        scroll = ctk.CTkScrollableFrame(
            popup, fg_color=AppTheme.BG_CARD, corner_radius=12,
            scrollbar_button_color=AppTheme.BG_INPUT)
        scroll.pack(fill="both", expand=True, padx=16, pady=(0, 16))

        def _selecionar(texto: str):
            self.motivo.delete("1.0", END)
            self.motivo.insert("1.0", texto)
            popup.destroy()

        for mot in motivos:
            row = ctk.CTkFrame(scroll, fg_color="transparent")
            row.pack(fill="x", pady=2)

            lbl = ctk.CTkLabel(
                row, text=mot,
                font=("Segoe UI", 12),
                text_color=AppTheme.TXT_MAIN,
                anchor="w", wraplength=440, justify="left",
                cursor="hand2",
            )
            lbl.pack(fill="x", padx=14, pady=6)
            lbl.bind("<Enter>",  lambda e, l=lbl: l.configure(text_color=_VERDE))
            lbl.bind("<Leave>",  lambda e, l=lbl: l.configure(text_color=AppTheme.TXT_MAIN))
            lbl.bind("<Button-1>", lambda e, m=mot: _selecionar(m))

    def _on_close(self):
        self.controller.fechar_driver()
        if hasattr(self.master, "destroy"):
            self.master.destroy()

    # ────────────────────────────────────────────────────────────────
    # ABA 1 — ANÁLISE
    # ────────────────────────────────────────────────────────────────
    def _build_analise(self):
        wrap = ctk.CTkFrame(self.tab_analise, fg_color="transparent")
        wrap.pack(fill="both", expand=True, padx=32, pady=24)

        # cabeçalho
        hdr = ctk.CTkFrame(wrap, fg_color="transparent")
        hdr.pack(fill="x", pady=(0, 20))
        ctk.CTkLabel(hdr, text="Análise de Processos",
                     font=("Segoe UI", 24, "bold"),
                     text_color=AppTheme.TXT_MAIN).pack(side="left")

        # barra de ações
        bar = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                           corner_radius=14)
        bar.pack(fill="x", pady=(0, 16))

        ctk.CTkCheckBox(
            bar, text="URGÊNCIA",
            variable=self.var_urgencia,
            fg_color=_AMBER, hover_color="#d97706",
            text_color=AppTheme.TXT_MAIN,
            font=("Segoe UI", 12, "bold")
        ).pack(side="left", padx=20, pady=14)

        # separador vertical
        ctk.CTkFrame(bar, width=1, height=28,
                     fg_color=AppTheme.BG_INPUT).pack(side="left", pady=10)

        for txt, cmd, cor, hov in [
            ("  Selecionar PDFs",          self.select_pdfs,          _VERDE,    _VERDE_H),
            ("  SEFAZ + Google Maps",      self._abrir_sefaz_thread,  _AZUL,     _AZUL_H),
            ("  Atualizar",                self.carregar_do_banco,    _MUTED,    AppTheme.BG_INPUT),
        ]:
            ctk.CTkButton(bar, text=txt, height=36, corner_radius=10,
                          fg_color=cor, hover_color=hov,
                          font=("Segoe UI", 12),
                          text_color="#ffffff",
                          command=cmd).pack(side="left", padx=(12, 0), pady=14)

        ctk.CTkButton(bar, text="  Histórico",
                      height=36, corner_radius=10,
                      fg_color=AppTheme.BG_INPUT,
                      hover_color=AppTheme.BG_APP,
                      font=("Segoe UI", 12),
                      text_color=AppTheme.TXT_MAIN,
                      command=self.abrir_historico
                      ).pack(side="right", padx=20, pady=14)

        ctk.CTkButton(bar, text="📄  Gerar Documento",
                      height=36, corner_radius=10,
                      fg_color="#6366f1", hover_color="#4f46e5",
                      font=("Segoe UI", 12),
                      text_color="#ffffff",
                      command=self._gerar_documento
                      ).pack(side="right", padx=(0, 8), pady=14)

        # ── Barra de filtro: memorando + ano ─────────────────────────────────
        filtro_bar = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                  corner_radius=14)
        filtro_bar.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(filtro_bar, text="🔍",
                     font=("Segoe UI", 14),
                     text_color=_MUTED).pack(side="left", padx=(16, 6), pady=12)

        self._filtro_memo_var = ctk.StringVar()
        self._filtro_memo_var.trace_add("write", lambda *_: self._aplicar_filtro())
        ctk.CTkEntry(
            filtro_bar, textvariable=self._filtro_memo_var,
            placeholder_text="Filtrar por nº do memorando (ex: 001)",
            height=34, width=220, corner_radius=8,
            fg_color=AppTheme.BG_INPUT, border_color=AppTheme.BG_INPUT,
            font=("Segoe UI", 12), text_color=AppTheme.TXT_MAIN,
        ).pack(side="left", pady=12)

        ctk.CTkLabel(filtro_bar, text="/",
                     font=("Segoe UI", 16, "bold"),
                     text_color=_MUTED).pack(side="left", padx=4)

        from datetime import datetime as _dt
        self._filtro_ano_var = ctk.StringVar(value=str(_dt.now().year))
        self._filtro_ano_var.trace_add("write", lambda *_: self._aplicar_filtro())
        ctk.CTkEntry(
            filtro_bar, textvariable=self._filtro_ano_var,
            placeholder_text="Ano (ex: 2025)",
            height=34, width=90, corner_radius=8,
            fg_color=AppTheme.BG_INPUT, border_color=AppTheme.BG_INPUT,
            font=("Segoe UI", 12), text_color=AppTheme.TXT_MAIN,
        ).pack(side="left", pady=12)

        ctk.CTkButton(
            filtro_bar, text="Limpar filtro", width=110, height=34,
            corner_radius=8,
            fg_color=AppTheme.BG_INPUT, hover_color=AppTheme.BG_APP,
            font=("Segoe UI", 11), text_color=_MUTED,
            command=self._limpar_filtro,
        ).pack(side="left", padx=(10, 0), pady=12)

        self._lbl_filtro_total = ctk.CTkLabel(
            filtro_bar, text="",
            font=("Segoe UI", 11), text_color=_MUTED)
        self._lbl_filtro_total.pack(side="right", padx=16, pady=12)

        # cabeçalho da tabela
        thead = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_INPUT,
                             corner_radius=10)
        thead.pack(fill="x", pady=(0, 4))
        for txt, w, col in [
            ("",          40,  0),
            ("Nome do PDF", 0, 1),
            ("Prioridade", 100, 2),
            ("Data/Hora",  140, 3),
            ("Ações",       90, 4),
        ]:
            thead.grid_columnconfigure(col, weight=(1 if col == 1 else 0),
                                       minsize=w)
            ctk.CTkLabel(thead, text=txt,
                         font=("Segoe UI", 11, "bold"),
                         text_color=_MUTED,
                         anchor="w"
                         ).grid(row=0, column=col,
                                padx=(16 if col == 0 else 8), pady=10,
                                sticky="w")

        # corpo da tabela
        self.tabela = ctk.CTkScrollableFrame(
            wrap, fg_color=AppTheme.BG_CARD, corner_radius=14,
            scrollbar_button_color=AppTheme.BG_INPUT)
        self.tabela.pack(fill="both", expand=True)

        # rodapé de ações em lote
        rodape = ctk.CTkFrame(wrap, fg_color="transparent")
        rodape.pack(fill="x", pady=(12, 0))

        ctk.CTkButton(rodape, text="Selecionar Todos",
                      height=40, corner_radius=10, width=140,
                      fg_color=AppTheme.BG_INPUT,
                      hover_color=AppTheme.BG_CARD,
                      font=("Segoe UI", 11),
                      text_color=AppTheme.TXT_MAIN,
                      command=self._selecionar_todos
                      ).pack(side="left")

        for txt, cmd, cor, hov in [
            ("Renovação",  lambda: self.processar("RENOVACAO"), _VERDE,    _VERDE_H),
            ("Inscrição",  lambda: self.processar("INSCRICAO"), _AZUL,     _AZUL_H),
            ("Devolução",  self.processar_devolucao,            _VERMELHO, _VERM_H),
            ("Enviar E-mail", self.abrir_popup_email,           _AMBER,    "#d97706"),
        ]:
            ctk.CTkButton(rodape, text=txt,
                          height=40, corner_radius=10,
                          fg_color=cor, hover_color=hov,
                          font=("Segoe UI", 12, "bold"),
                          text_color="#ffffff",
                          command=cmd).pack(side="left", padx=(10, 0))

    # ────────────────────────────────────────────────────────────────
    # ABA 2 — ADICIONAR
    # ────────────────────────────────────────────────────────────────
    def _build_adicionar(self):
        outer = ctk.CTkScrollableFrame(
            self.tab_adicionar, fg_color="transparent",
            scrollbar_button_color=AppTheme.BG_INPUT)
        outer.pack(fill="both", expand=True)

        wrap = ctk.CTkFrame(outer, fg_color="transparent")
        wrap.pack(fill="both", expand=True, padx=40, pady=28)

        ctk.CTkLabel(wrap, text="Adicionar Registro",
                     font=("Segoe UI", 24, "bold"),
                     text_color=AppTheme.TXT_MAIN
                     ).pack(anchor="w", pady=(0, 20))

        # tipo selector
        tipo_card = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                 corner_radius=14)
        tipo_card.pack(fill="x", pady=(0, 20))

        self.radio = ctk.StringVar(value="insc")
        tipo_inner = ctk.CTkFrame(tipo_card, fg_color="transparent")
        tipo_inner.pack(pady=16)

        for txt, val in [("Inscrição", "insc"), ("Devolução", "dev")]:
            ctk.CTkRadioButton(
                tipo_inner, text=txt,
                variable=self.radio, value=val,
                command=self._reset_form,
                font=("Segoe UI", 13, "bold"),
                fg_color=_VERDE, hover_color=_VERDE_H,
                text_color=AppTheme.TXT_MAIN,
            ).pack(side="left", padx=28)

        # formulário
        form_card = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                 corner_radius=14)
        form_card.pack(fill="x")
        self.form = ctk.CTkFrame(form_card, fg_color="transparent")
        self.form.pack(fill="both", expand=True, padx=28, pady=24)

        def _entry(placeholder):
            return ctk.CTkEntry(
                self.form, placeholder_text=placeholder,
                height=44, corner_radius=12,
                font=("Segoe UI", 13),
                fg_color=AppTheme.BG_INPUT,
                border_color=AppTheme.BG_INPUT,
                text_color=AppTheme.TXT_MAIN,
            )

        def _label(txt):
            ctk.CTkLabel(self.form, text=txt,
                         font=("Segoe UI", 11, "bold"),
                         text_color=_MUTED,
                         anchor="w").pack(fill="x", pady=(10, 3))

        # Campos comuns
        _label("Nome Completo *")
        self.nome = _entry("Ex: João da Silva")
        self.nome.pack(fill="x")

        _label("CPF *")
        self.cpf = _entry("000.000.000-00")
        self.cpf.pack(fill="x")
        self.cpf.bind("<KeyRelease>", self._formatar_cpf)

        _label("Município *")
        self.municipio = _entry("Ex: Manaus")
        self.municipio.pack(fill="x")

        _label("Memorando *")
        self.memorando = _entry("Ex: 001/2025")
        self.memorando.pack(fill="x")

        # Campo exclusivo inscrição
        _label("Tipo")
        self._lbl_tipo = self.form.winfo_children()[-1]
        self.tipo = ctk.CTkComboBox(
            self.form, values=["INSC", "RENOV"],
            height=44, corner_radius=12,
            font=("Segoe UI", 13),
            fg_color=AppTheme.BG_INPUT,
            border_color=AppTheme.BG_INPUT,
            button_color=_VERDE,
            button_hover_color=_VERDE_H,
            text_color=AppTheme.TXT_MAIN,
            dropdown_fg_color=AppTheme.BG_CARD,
            dropdown_text_color=AppTheme.TXT_MAIN,
        )
        self.tipo.pack(fill="x")

        # Campos exclusivos devolução (ocultos inicialmente)
        self._lbl_categoria = ctk.CTkLabel(
            self.form, text="Categoria *",
            font=("Segoe UI", 11, "bold"),
            text_color=_MUTED, anchor="w")

        self.categoria = ctk.CTkComboBox(
            self.form,
            values=["Endereço", "Documentos", "Cadastro",
                    "Pesca", "Simples Nacional", "Animais", "Outros"],
            height=44, corner_radius=12,
            font=("Segoe UI", 13),
            fg_color=AppTheme.BG_INPUT,
            border_color=AppTheme.BG_INPUT,
            button_color=_VERDE,
            button_hover_color=_VERDE_H,
            text_color=AppTheme.TXT_MAIN,
            dropdown_fg_color=AppTheme.BG_CARD,
            dropdown_text_color=AppTheme.TXT_MAIN,
            command=self._abrir_motivos,
        )

        self._lbl_motivo = ctk.CTkLabel(
            self.form, text="Motivo da Devolução *",
            font=("Segoe UI", 11, "bold"),
            text_color=_MUTED, anchor="w")

        self.motivo = ctk.CTkTextbox(
            self.form, height=100, corner_radius=12,
            font=("Segoe UI", 13),
            fg_color=AppTheme.BG_INPUT,
            border_color=AppTheme.BG_INPUT,
            text_color=AppTheme.TXT_MAIN,
            border_width=0,
        )

        # Botões lado a lado — dentro do form para reposicionar no _reset_form
        self._btn_row = ctk.CTkFrame(self.form, fg_color="transparent")
        self._btn_row.grid_columnconfigure((0, 1), weight=1)

        self._btn_salvar = ctk.CTkButton(
            self._btn_row, text="Salvar Registro",
            height=50, corner_radius=14,
            font=("Segoe UI", 14, "bold"),
            fg_color=_VERDE, hover_color=_VERDE_H,
            text_color="#ffffff",
            command=self._salvar_registro,
        )
        self._btn_salvar.grid(row=0, column=0, padx=(0, 6), sticky="ew")

        self._btn_email_form = ctk.CTkButton(
            self._btn_row, text="Enviar E-mail",
            height=50, corner_radius=14,
            font=("Segoe UI", 14, "bold"),
            fg_color=_AMBER, hover_color="#d97706",
            text_color="#ffffff",
            command=self._enviar_email_form,
        )
        self._btn_email_form.grid(row=0, column=1, padx=(6, 0), sticky="ew")

        self._btn_row.pack(fill="x", pady=(20, 0))

    # ────────────────────────────────────────────────────────────────
    # Linha da tabela
    # ────────────────────────────────────────────────────────────────
    def _criar_linha(self, reg: dict):
        urgente = reg.get("urgencia", False)
        # Cor de fundo sutil para urgentes
        bg = "#1c1007" if urgente else AppTheme.BG_CARD

        linha = ctk.CTkFrame(self.tabela, fg_color=bg,
                             corner_radius=10)
        linha.pack(fill="x", pady=3, padx=6)
        linha.grid_columnconfigure(1, weight=1)

        var = ctk.BooleanVar()
        self.registros_vars[reg["id"]] = var

        ctk.CTkCheckBox(
            linha, variable=var, text="", width=40,
            fg_color=_VERDE, hover_color=_VERDE_H,
        ).grid(row=0, column=0, padx=(14, 6), pady=12)

        # Nome clicável
        lbl = ctk.CTkLabel(
            linha, text=reg["nome_pdf"],
            anchor="w", cursor="hand2",
            font=("Segoe UI", 12),
            text_color=_AZUL,
        )
        lbl.grid(row=0, column=1, padx=6, pady=12, sticky="ew")
        lbl.bind("<Button-1>",
                 lambda _, aid=reg["id"]: self.visualizar_pdf(aid))

        # Badge urgência
        if urgente:
            _badge(linha, "URGENTE", _AMBER
                   ).grid(row=0, column=2, padx=6, pady=12)
        else:
            ctk.CTkLabel(linha, text="Normal", width=90,
                         font=("Segoe UI", 11),
                         text_color=_MUTED,
                         ).grid(row=0, column=2, padx=6, pady=12)

        # Data
        criado = reg.get("criado_em")
        if isinstance(criado, str):
            try:
                criado = datetime.fromisoformat(criado)
            except ValueError:
                criado = datetime.now()
        data_txt = criado.strftime("%d/%m/%Y %H:%M") if criado else "—"
        ctk.CTkLabel(linha, text=data_txt, width=130,
                     font=("Segoe UI", 11), text_color=_MUTED,
                     ).grid(row=0, column=3, padx=6, pady=12)

        # Botões
        btns = ctk.CTkFrame(linha, fg_color="transparent")
        btns.grid(row=0, column=4, padx=(6, 14), pady=12)

        ctk.CTkButton(
            btns, text="Ver", width=44, height=30,
            corner_radius=8,
            fg_color=_AZUL, hover_color=_AZUL_H,
            font=("Segoe UI", 11, "bold"), text_color="#fff",
            command=lambda aid=reg["id"]: self.visualizar_pdf(aid)
        ).pack(side="left", padx=2)

        ctk.CTkButton(
            btns, text="↓", width=32, height=30,
            corner_radius=8,
            fg_color=AppTheme.BG_INPUT,
            hover_color=AppTheme.BG_APP,
            font=("Segoe UI", 12), text_color=AppTheme.TXT_MAIN,
            command=lambda aid=reg["id"],
                           n=reg["nome_pdf"]: self.baixar_pdf(aid, n)
        ).pack(side="left", padx=2)

    # ────────────────────────────────────────────────────────────────
    # Helpers — reset form
    # ────────────────────────────────────────────────────────────────
    def _reset_form(self):
        """Mostra/oculta campos conforme tipo selecionado."""
        self._btn_row.pack_forget()

        if self.radio.get() == "dev":
            self._lbl_tipo.pack_forget()
            self.tipo.pack_forget()
            self._lbl_categoria.pack(fill="x", pady=(10, 3))
            self.categoria.pack(fill="x")
            self._lbl_motivo.pack(fill="x", pady=(10, 3))
            self.motivo.pack(fill="x")
        else:
            self._lbl_categoria.pack_forget()
            self.categoria.pack_forget()
            self._lbl_motivo.pack_forget()
            self.motivo.pack_forget()
            self._lbl_tipo.pack(fill="x", pady=(10, 3))
            self.tipo.pack(fill="x")

        self._btn_row.pack(fill="x", pady=(20, 0))

    # ────────────────────────────────────────────────────────────────
    # Formatação CPF — BUG CORRIGIDO (cursor não salta mais)
    # ────────────────────────────────────────────────────────────────
    def _formatar_cpf(self, e=None):
        texto = self.cpf.get()
        # Salvar posição relativa aos dígitos ANTES
        pos_raw = self.cpf.index("insert")
        digitos_antes = sum(1 for c in texto[:pos_raw] if c.isdigit())

        formatado = self.controller.formatar_cpf(texto)
        if formatado == texto:
            return

        self.cpf.delete(0, END)
        self.cpf.insert(0, formatado)

        # Reposicionar cursor: avançar pelos dígitos equivalentes
        nova_pos = 0
        contados = 0
        for i, c in enumerate(formatado):
            if contados >= digitos_antes:
                nova_pos = i
                break
            if c.isdigit():
                contados += 1
        else:
            nova_pos = len(formatado)

        self.cpf.icursor(nova_pos)

    # ────────────────────────────────────────────────────────────────
    # Carregar registros — em thread para não travar UI
    # ────────────────────────────────────────────────────────────────
    def _carregar_thread(self):
        try:
            registros = self.controller.carregar_registros_pendentes()
            self.after(0, self._preencher_tabela, registros)
        except Exception as exc:
            self.after(0, messagebox.showerror,
                       "Erro", f"Erro ao carregar dados:\n{exc}")

    def carregar_do_banco(self):
        """Limpa tabela e recarrega em thread."""
        for w in self.tabela.winfo_children():
            w.destroy()
        self.registros_vars.clear()
        ctk.CTkLabel(self.tabela, text="Carregando...",
                     text_color=_MUTED,
                     font=("Segoe UI", 12)
                     ).pack(pady=30)
        threading.Thread(target=self._carregar_thread, daemon=True).start()

    def _preencher_tabela(self, registros: list):
        self._todos_registros = registros   # cache para filtragem local
        self._aplicar_filtro()

    def _limpar_filtro(self):
        self._filtro_memo_var.set("")
        self._filtro_ano_var.set("")

    def _aplicar_filtro(self):
        import re as _re
        memo = self._filtro_memo_var.get().strip().lower()
        ano  = self._filtro_ano_var.get().strip()

        registros = getattr(self, "_todos_registros", [])
        if memo or ano:
            filtrados = []
            for r in registros:
                nome   = (r.get("nome_pdf")   or "").lower()
                memo_r = (r.get("memorando")  or "").lower()
                match  = _re.search(r'(\d+)[/\-](\d{4})', nome + " " + memo_r)
                num_r  = match.group(1) if match else ""
                ano_r  = match.group(2) if match else ""
                ok_memo = (not memo) or memo in num_r or memo in memo_r
                ok_ano  = (not ano)  or ano == ano_r or ano in nome
                if ok_memo and ok_ano:
                    filtrados.append(r)
        else:
            filtrados = registros

        for w in self.tabela.winfo_children():
            w.destroy()
        self.registros_vars.clear()

        total    = len(registros)
        visiveis = len(filtrados)
        try:
            txt = f"{visiveis} de {total}" if (memo or ano) else f"{total} registro(s)"
            self._lbl_filtro_total.configure(text=txt)
        except Exception:
            pass

        if not filtrados:
            msg = ("Nenhum registro encontrado com esse filtro."
                   if (memo or ano) else "Nenhum registro pendente.")
            ctk.CTkLabel(self.tabela, text=msg,
                         text_color=_MUTED,
                         font=("Segoe UI", 13)).pack(pady=40)
            return
        for reg in filtrados:
            self._criar_linha(reg)

    def _gerar_documento(self):
        """Abre diálogo para salvar e gera o .docx de unidades e motivos."""
        import subprocess, sys, os, tempfile, shutil
        caminho = filedialog.asksaveasfilename(
            title="Salvar documento",
            defaultextension=".docx",
            initialfile="Unidades_e_Motivos_Devolucao.docx",
            filetypes=[("Word", "*.docx"), ("Todos", "*.*")],
        )
        if not caminho:
            return
        try:
            # Usa python-docx para geração simples sem Node.js em runtime
            from docx import Document as DocxDoc
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = DocxDoc()
            style = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(11)

            # Título
            t = doc.add_heading("IDAM — CPP · Unidades Locais e Motivos de Devolução", 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Seção 1
            doc.add_heading("1. Unidades Locais — Contatos de E-mail", 1)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "MUNICÍPIO"
            hdr[1].text = "E-MAIL"
            for cell in hdr:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

            for mun, email in sorted(self._EMAILS_UNIDADES.items()):
                row = table.add_row().cells
                row[0].text = mun
                row[1].text = email

            doc.add_page_break()

            # Seção 2
            doc.add_heading("2. Motivos de Devolução por Categoria", 1)
            for cat, motivos in self._motivos.items():
                doc.add_heading(cat, 2)
                if not motivos:
                    p = doc.add_paragraph("Sem motivos cadastrados — preencher manualmente.")
                    p.italic = True
                else:
                    for m in motivos:
                        doc.add_paragraph(m, style="List Bullet")

            doc.save(caminho)
            messagebox.showinfo("Documento gerado",
                                f"Arquivo salvo em:\n{caminho}")
        except ImportError:
            messagebox.showerror(
                "Dependência ausente",
                "Instale python-docx:\n  pip install python-docx"
            )
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao gerar documento:\n{e}")

    # ────────────────────────────────────────────────────────────────
    # Selecionar todos
    # ────────────────────────────────────────────────────────────────
    def _selecionar_todos(self):
        if not self.registros_vars:
            return
        novo = not all(v.get() for v in self.registros_vars.values())
        for v in self.registros_vars.values():
            v.set(novo)

    # ────────────────────────────────────────────────────────────────
    # Ações
    # ────────────────────────────────────────────────────────────────
    def select_pdfs(self):
        caminhos = filedialog.askopenfilenames(
            title="Selecionar PDFs",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if not caminhos:
            return
        sucesso, erros = self.controller.salvar_pdfs(
            caminhos, self.var_urgencia.get())
        msg = f"{sucesso} PDF(s) adicionado(s)."
        if erros:
            msg += f"\n\nErros ({len(erros)}):\n" + "\n".join(erros[:5])
            if len(erros) > 5:
                msg += f"\n... e mais {len(erros) - 5}."
        if sucesso:
            messagebox.showinfo("Sucesso", msg)
            self.carregar_do_banco()
        else:
            messagebox.showerror("Erro", msg)

    def visualizar_pdf(self, analise_id):
        ok, msg = self.controller.visualizar_pdf(analise_id)
        if not ok:
            messagebox.showerror("Erro", msg)

    def baixar_pdf(self, analise_id, nome):
        ok, msg = self.controller.baixar_pdf(analise_id, nome)
        if ok:
            messagebox.showinfo("Salvo", f"PDF salvo em:\n{msg}")
        elif msg != "Operação cancelada":
            messagebox.showerror("Erro", msg)

    def processar(self, novo_status: str):
        ids = [i for i, v in self.registros_vars.items() if v.get()]
        if not ids:
            messagebox.showwarning("Atenção", "Nenhum registro selecionado.")
            return
        ok, msg = self.controller.processar_registros(ids, novo_status)
        if ok:
            messagebox.showinfo("Sucesso", msg)
            self.carregar_do_banco()
        else:
            messagebox.showerror("Erro", msg)

    def processar_devolucao(self):
        ids = [i for i, v in self.registros_vars.items() if v.get()]
        if not ids:
            messagebox.showwarning("Atenção", "Nenhum registro selecionado.")
            return

        def _confirmar(motivo):
            """Chamado pelo DevolucaoPopup após escolher o motivo."""
            # Busca dados do primeiro registro para pré-preencher o formulário
            primeiro = {}
            try:
                primeiro = self.controller.repo.buscar_por_id(ids[0]) or {}
            except Exception:
                pass
            self._abrir_popup_dados_devolucao(ids, motivo, primeiro)

        DevolucaoPopup(self, _confirmar)

    def _abrir_popup_dados_devolucao(self, ids: list, motivo: str, dados_pre: dict):
        """
        Segundo passo da devolução: coleta nome, CPF, município e memorando
        para salvar no banco separado (analiseap) e marca no banco principal.
        """
        popup = ctk.CTkToplevel(self)
        popup.title("Registrar Devolução — Dados Complementares")
        popup.geometry("520x560")
        popup.resizable(False, False)
        popup.configure(fg_color=AppTheme.BG_APP)
        popup.grab_set()
        popup.after(0, lambda: popup.geometry(
            f"520x560+"
            f"{popup.winfo_screenwidth()  // 2 - 260}+"
            f"{popup.winfo_screenheight() // 2 - 280}"
        ))

        wrap = ctk.CTkFrame(popup, fg_color="transparent")
        wrap.pack(fill="both", expand=True, padx=28, pady=24)

        # Cabeçalho
        ctk.CTkLabel(wrap, text="Dados da Devolução",
                     font=("Segoe UI", 18, "bold"),
                     text_color=AppTheme.TXT_MAIN).pack(anchor="w", pady=(0, 4))
        motivo_resumo = motivo[:60] + "…" if len(motivo) > 60 else motivo
        ctk.CTkLabel(wrap, text=f"Motivo: {motivo_resumo}",
                     font=("Segoe UI", 11), text_color=_MUTED
                     ).pack(anchor="w", pady=(0, 16))

        # Formulário
        card = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD, corner_radius=14)
        card.pack(fill="x", pady=(0, 20))
        form = ctk.CTkFrame(card, fg_color="transparent")
        form.pack(fill="x", padx=20, pady=20)

        def _entry(placeholder, valor=""):
            e = ctk.CTkEntry(
                form, placeholder_text=placeholder,
                height=42, corner_radius=10,
                font=("Segoe UI", 12),
                fg_color=AppTheme.BG_INPUT,
                border_color=AppTheme.BG_INPUT,
                text_color=AppTheme.TXT_MAIN,
            )
            if valor:
                e.insert(0, valor)
            return e

        def _label(txt):
            ctk.CTkLabel(form, text=txt,
                         font=("Segoe UI", 11, "bold"),
                         text_color=_MUTED, anchor="w").pack(fill="x", pady=(8, 2))

        # Pré-preenche com dados do banco quando disponíveis
        cpf_pre  = dados_pre.get("cpf", "")
        if cpf_pre and len(cpf_pre) == 11:
            cpf_pre = f"{cpf_pre[:3]}.{cpf_pre[3:6]}.{cpf_pre[6:9]}-{cpf_pre[9:]}"
        memo_pre = dados_pre.get("memorando", "")
        muni_pre = dados_pre.get("unloc", "") or dados_pre.get("municipio", "")

        _label("Nome Completo *")
        entry_nome = _entry("Ex: João da Silva")
        entry_nome.pack(fill="x")

        _label("CPF *")
        entry_cpf = _entry("000.000.000-00", cpf_pre)
        entry_cpf.pack(fill="x")

        _label("Município *")
        entry_muni = _entry("Ex: Manaus", muni_pre)
        entry_muni.pack(fill="x")

        _label("Memorando *")
        entry_memo = _entry("Ex: 001/2025", memo_pre)
        entry_memo.pack(fill="x")

        # Botões
        btns = ctk.CTkFrame(wrap, fg_color="transparent")
        btns.pack(fill="x")
        btns.grid_columnconfigure((0, 1), weight=1)

        def _salvar():
            nome = entry_nome.get().strip()
            cpf  = entry_cpf.get().strip()
            muni = entry_muni.get().strip()
            memo = entry_memo.get().strip()

            if not nome:
                messagebox.showwarning("Atenção", "Informe o nome completo!", parent=popup)
                entry_nome.focus(); return
            if not self.controller.validar_cpf(cpf):
                messagebox.showwarning("Atenção", "CPF deve ter 11 dígitos!", parent=popup)
                entry_cpf.focus(); return
            if not muni:
                messagebox.showwarning("Atenção", "Informe o município!", parent=popup)
                entry_muni.focus(); return
            if not memo:
                messagebox.showwarning("Atenção", "Informe o memorando!", parent=popup)
                entry_memo.focus(); return

            # Extrai categoria do motivo (formato: "[Categoria] detalhe")
            import re as _re
            cat_match = _re.match(r'\[([^\]]+)\]', motivo)
            categoria = cat_match.group(1) if cat_match else "Outros"

            # 1. Marca no banco principal como DEVOLUCAO
            ok, msg = self.controller.processar_devolucao(ids, motivo)

            # 2. Salva no banco separado (analiseap)
            aviso_extra = ""
            try:
                ok2, msg2 = self.controller.salvar_registro("dev", {
                    "nome":      nome,
                    "cpf":       self.controller.limpar_cpf(cpf),
                    "municipio": muni,
                    "memorando": memo,
                    "categoria": categoria,
                    "motivo":    motivo,
                })
                if not ok2:
                    aviso_extra = f"\n\n⚠️ Banco separado não salvo:\n{msg2}"
            except Exception as exc:
                aviso_extra = f"\n\n⚠️ Banco separado não salvo:\n{exc}"

            popup.destroy()

            if ok:
                messagebox.showinfo("Sucesso", msg + aviso_extra)
                self.carregar_do_banco()
                if callable(self._on_refresh_devolucao):
                    try:
                        self._on_refresh_devolucao()
                    except Exception:
                        pass
            else:
                messagebox.showerror("Erro", msg)

        ctk.CTkButton(
            btns, text="Confirmar Devolução",
            height=46, corner_radius=12,
            fg_color=_VERMELHO, hover_color=_VERM_H,
            font=("Segoe UI", 13, "bold"), text_color="#fff",
            command=_salvar,
        ).grid(row=0, column=0, padx=(0, 6), sticky="ew")

        ctk.CTkButton(
            btns, text="Cancelar",
            height=46, corner_radius=12,
            fg_color=AppTheme.BG_INPUT, hover_color=AppTheme.BG_APP,
            font=("Segoe UI", 13), text_color=AppTheme.TXT_MAIN,
            command=popup.destroy,
        ).grid(row=0, column=1, padx=(6, 0), sticky="ew")

    def _abrir_sefaz_thread(self):
        threading.Thread(target=self._abrir_sefaz, daemon=True).start()

    def _abrir_sefaz(self):
        ok, msg = self.controller.abrir_sefaz()
        if not ok:
            self.after(0, messagebox.showerror, "Erro", msg)

    def _salvar_registro(self):
        nome     = self.nome.get().strip()
        cpf_txt  = self.cpf.get().strip()
        muni     = self.municipio.get().strip()
        memo_val = self.memorando.get().strip()

        # Validações
        if not nome:
            messagebox.showwarning("Atenção", "Informe o nome completo!")
            self.nome.focus(); return
        if not self.controller.validar_cpf(cpf_txt):
            messagebox.showwarning("Atenção", "CPF deve conter 11 dígitos!")
            self.cpf.focus(); return
        if not muni:
            messagebox.showwarning("Atenção", "Informe o município!")
            self.municipio.focus(); return
        if not memo_val:
            messagebox.showwarning("Atenção", "Informe o memorando!")
            self.memorando.focus(); return

        tipo_reg = self.radio.get()
        dados = {
            "nome":       nome,
            "cpf":        self.controller.limpar_cpf(cpf_txt),
            "municipio":  muni,
            "memorando":  memo_val,
        }

        if tipo_reg == "insc":
            if not self.tipo.get():
                messagebox.showwarning("Atenção", "Selecione o tipo!"); return
            dados["tipo"] = self.tipo.get()
        else:
            if not self.categoria.get():
                messagebox.showwarning("Atenção", "Selecione a categoria!"); return
            motivo_txt = self.motivo.get("1.0", END).strip()
            if not motivo_txt:
                messagebox.showwarning("Atenção", "Informe o motivo!"); return
            dados["categoria"] = self.categoria.get()
            dados["motivo"]    = motivo_txt

        ok, msg = self.controller.salvar_registro(tipo_reg, dados)
        if ok:
            # Limpar campos comuns
            for w in (self.nome, self.cpf, self.municipio, self.memorando):
                w.delete(0, END)
            # Limpar motivo apenas quando estava visível (modo devolução)
            if tipo_reg == "dev":
                self.motivo.delete("1.0", END)
            self.tipo.set("INSC")
            self.categoria.set("Endereço")
            messagebox.showinfo("Sucesso", msg)
        else:
            messagebox.showerror("Erro", msg)

    # ────────────────────────────────────────────────────────────────
    # E-mail para unidade local
    # ────────────────────────────────────────────────────────────────
    _EMAILS_UNIDADES = {
        "TESTE TESTE": "carteiradoprodutor@gmail.com",
        "ALVARÃES": "unlocalvaraes@idam.am.gov.br",
        "AMATURÁ": "unlocamatura@idam.am.gov.br",
        "ANAMÃ": "unlocanama@idam.am.gov.br",
        "ANORI": "euricopaulobarbosa@gmail.com",
        "BALBINA": "unlocbalbina@idam.am.gov.br",
        "APUÍ": "idamapui@hotmail.com",
        "ATALAIA DO NORTE": "idamatalaia@gmail.com",
        "AUTAZES": "unlocautazes@gmail.com",
        "BARCELOS": "unlocbarcelos@idam.am.gov.br",
        "BARREIRINHA": "unlocbarreirinha@idam.am.gov.br",
        "BENJAMIN CONSTANT": "unlocbenjaminconstant@idam.am.gov.br",
        "BERURI": "unlocberuri@idam.am.gov.br",
        "BOA VISTA DO RAMOS": "arlindobvr@gmail.com",
        "BOCA DO ACRE": "idambocadoacre@hotmail.com",
        "BORBA": "aguinaldoceplac@gmail.com",
        "CAAPIRANGA": "unloccaapiranga@idam.am.gov.br",
        "CANUTAMA": "batistagdau@gmail.com",
        "CARAUARI": "unloccarauari@idam.am.gov.br",
        "CAREIRO DA VÁRZEA": "cppcareirovarzea@gmail.com",
        "CAREIRO": "unloccareiro@gmail.com",
        "COARI": "unloccoari@idam.am.gov.br",
        "CODAJÁS": "unloccodajas@idam.am.gov.br",
        "EIRUNEPÉ": "unloceirunepe@idam.am.gov.br",
        "ENVIRA": "unlocenvira@gmail.com",
        "FONTE BOA": "unlocfonteboa@idam.am.gov.br",
        "GUAJARÁ": "idamguajara@gmail.com",
        "HUMAITÁ": "unlochumaita@idam.am.gov.br",
        "IPIXUNA": "UNLOCIPIXUNA@idam.am.gov.br",
        "IRANDUBA": "unlocirandubacpp@gmail.com",
        "ITACOATIARA": "unlocitacoatiara@idam.am.gov.br",
        "ITAMARATI": "unlocitamarati@idam.am.gov.br",
        "ITAPIRANGA": "unlocitapiranga@idam.am.gov.br",
        "JAPURÁ": "unlocjapura@idam.Am.gov.br",
        "JURUÁ": "ffsilva20@gmail.com",
        "JUTAÍ": "ramosampaio@hotmail.com",
        "LÁBREA": "unloclabrea@idam.am.gov.br",
        "MANACAPURU": "unlocmanacapuru@idam.am.gov.br",
        "MANAQUIRI": "unlocmanaquiri@idam.am.gov.br",
        "MANAUS": "unlocmanaus@gmail.com",
        "MANAUS ZONA LESTE": "idamzonaleste@gmail.com",
        "MANICORÉ": "unlocmanicore@idam.am.gov.br",
        "MARAÃ": "deividdelrikmaraa@gmail.com",
        "MAUÉS": "unlocmaues@idam.am.gov.br",
        "NHAMUNDÁ": "unlocnhamunda@idam.am.gov.br",
        "NOVA OLINDA DO NORTE": "unlocnovaolindanorte@idam.am.gov.br",
        "NOVO AIRÃO": "unlocnovoairao@idam.am.gov.br",
        "NOVO ARIPUANÃ": "unlocnovoaripuana@idam.am.gov.br",
        "NOVO REMANSO": "lucianolobovaz@gmail.com",
        "SANTO ANTÔNIO DO MATUPI": "unlocmatupi@idam.am.gov.br",
        "PARINTINS": "idamparintins@gmail.com",
        "PAUINI": "simoneandrade21@hotmail.com",
        "PRESIDENTE FIGUEIREDO": "unlocpresidentefigueiredo@gmail.com",
        "RIO PRETO DA EVA": "unlocriopretoeva@idam.am.gov.br",
        "SANTA ISABEL DO RIO NEGRO": "unlocstaizabelrionegro@idam.am.gov.br",
        "SANTO ANTÔNIO DO IÇÁ": "unlocsantoantonioica@idam.am.gov.br",
        "SÃO GABRIEL DA CACHOEIRA": "unlocsaogabrielcachoeira@idam.am.gov.br",
        "SÃO PAULO DE OLIVENÇA": "unlocsaopauloolivenca@idam.am.gov.br",
        "SÃO SEBASTIÃO DO UATUMÃ": "idam.ssu@gmail.com",
        "SILVES": "unlocsilves@idam.am.gov.br",
        "TABATINGA": "zecoelhodavila@hotmail.com",
        "TAPAUÁ": "unloctapaua@idam.am.gov.br",
        "TEFÉ": "idamtefe@gmail.com",
        "TONANTINS": "unloctonantins@idam.am.gov.br",
        "UARINI": "unlocuarini@idam.am.gov.br",
        "URUCARÁ": "idam.urucara@gmail.com",
        "URUCURITUBA": "unlocurucurituba@idam.am.gov.br",
        "VILA EXTREMA": "unlocextrema@gmail.com",
        "VILA RICA DE CAVIANA": "unlocvilaricacaviana@idam.am.gov.br",
        "SUL DE CANUTAMA": "idamsuldecanutama@gmail.com",
        "PA AGROVILA DE CABURI/PARINTINS": "asaelrocha@hotmail.com",
        "MONTE SINAI": "unlocpamontesinai@idam.am.gov.br",
        "VILA DE LINDÓIA": "unloclindoia@idam.am.gov.br",
        "VILA DA REALIDADE": "unlocparealidade@idam.am.gov.br",
        "FOZ DE CANUMÃ": "vmsilvafilho@hotmail.com",
    }

    def abrir_popup_email(self):
        """Abre popup para selecionar unidade e enviar e-mail com os PDFs selecionados."""
        ids = [i for i, v in self.registros_vars.items() if v.get()]
        if not ids:
            messagebox.showwarning("Atenção", "Selecione ao menos um PDF para enviar.")
            return

        popup = ctk.CTkToplevel(self)
        popup.title("Enviar E-mail para Unidade Local")
        popup.geometry("560x620")
        popup.resizable(False, False)
        popup.configure(fg_color=AppTheme.BG_APP)
        popup.grab_set()
        popup.after(0, lambda: popup.geometry(
            f"560x620+"
            f"{popup.winfo_screenwidth()  // 2 - 280}+"
            f"{popup.winfo_screenheight() // 2 - 310}"
        ))

        wrap = ctk.CTkFrame(popup, fg_color="transparent")
        wrap.pack(fill="both", expand=True, padx=28, pady=24)

        ctk.CTkLabel(wrap, text="Enviar E-mail para Unidade Local",
                     font=("Segoe UI", 18, "bold"),
                     text_color=AppTheme.TXT_MAIN).pack(anchor="w", pady=(0, 20))

        # ── Remetente ────────────────────────────────────────────────
        card_rem = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD, corner_radius=14)
        card_rem.pack(fill="x", pady=(0, 12))

        ctk.CTkLabel(card_rem, text="Configuração do Remetente",
                     font=("Segoe UI", 12, "bold"),
                     text_color=_MUTED).pack(anchor="w", padx=20, pady=(14, 8))

        def _entry_card(parent, placeholder):
            return ctk.CTkEntry(
                parent, placeholder_text=placeholder,
                height=40, corner_radius=10,
                font=("Segoe UI", 12),
                fg_color=AppTheme.BG_INPUT,
                border_color=AppTheme.BG_INPUT,
                text_color=AppTheme.TXT_MAIN,
            )

        ctk.CTkLabel(card_rem, text="E-mail remetente:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_email_rem = _entry_card(card_rem, "seu.email@gmail.com")
        entry_email_rem.pack(fill="x", padx=20, pady=(2, 8))

        ctk.CTkLabel(card_rem, text="Senha do app (Gmail):",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_senha = _entry_card(card_rem, "Senha de app do Google")
        entry_senha.configure(show="*")
        entry_senha.pack(fill="x", padx=20, pady=(2, 14))

        # ── Destinatário ─────────────────────────────────────────────
        card_dest = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD, corner_radius=14)
        card_dest.pack(fill="x", pady=(0, 12))

        ctk.CTkLabel(card_dest, text="Unidade Local Destinatária",
                     font=("Segoe UI", 12, "bold"),
                     text_color=_MUTED).pack(anchor="w", padx=20, pady=(14, 8))

        municipios = sorted(self._EMAILS_UNIDADES.keys())
        combo_muni = ctk.CTkComboBox(
            card_dest,
            values=municipios,
            height=40, corner_radius=10,
            font=("Segoe UI", 12),
            fg_color=AppTheme.BG_INPUT,
            border_color=AppTheme.BG_INPUT,
            button_color=_VERDE,
            button_hover_color=_VERDE_H,
            text_color=AppTheme.TXT_MAIN,
            dropdown_fg_color=AppTheme.BG_CARD,
            dropdown_text_color=AppTheme.TXT_MAIN,
        )
        combo_muni.pack(fill="x", padx=20, pady=(0, 6))

        lbl_email_dest = ctk.CTkLabel(card_dest, text="",
                                      font=("Segoe UI", 11),
                                      text_color=_VERDE)
        lbl_email_dest.pack(anchor="w", padx=20, pady=(0, 14))

        def _on_muni_change(choice):
            email = self._EMAILS_UNIDADES.get(choice, "")
            lbl_email_dest.configure(text=f"  {email}")

        combo_muni.configure(command=_on_muni_change)

        # ── Assunto e mensagem ────────────────────────────────────────
        card_msg = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD, corner_radius=14)
        card_msg.pack(fill="x", pady=(0, 16))

        ctk.CTkLabel(card_msg, text="Assunto:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20, pady=(14, 2))
        entry_assunto = _entry_card(card_msg, "Ex: Devolução de processos — Manaus")
        entry_assunto.pack(fill="x", padx=20, pady=(0, 8))

        ctk.CTkLabel(card_msg, text="Mensagem (opcional):",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        txt_mensagem = ctk.CTkTextbox(
            card_msg, height=80, corner_radius=10,
            font=("Segoe UI", 12),
            fg_color=AppTheme.BG_INPUT,
            border_width=0,
            text_color=AppTheme.TXT_MAIN,
        )
        txt_mensagem.pack(fill="x", padx=20, pady=(2, 14))
        txt_mensagem.insert("1.0", f"Prezados,\n\nSegue(m) em anexo o(s) PDF(s) referente(s) ao(s) processo(s) selecionado(s).\n\nAtenciosamente,\n{self.usuario}")

        # ── Botões ────────────────────────────────────────────────────
        btns = ctk.CTkFrame(wrap, fg_color="transparent")
        btns.pack(fill="x")
        btns.grid_columnconfigure((0, 1), weight=1)

        def _confirmar():
            remetente = entry_email_rem.get().strip()
            senha     = entry_senha.get().strip()
            municipio = combo_muni.get().strip()
            assunto   = entry_assunto.get().strip()
            mensagem  = txt_mensagem.get("1.0", "end-1c").strip()

            if not remetente or not senha:
                messagebox.showwarning("Atenção", "Preencha o e-mail e a senha do remetente.",
                                       parent=popup)
                return
            if not municipio or municipio not in self._EMAILS_UNIDADES:
                messagebox.showwarning("Atenção", "Selecione uma unidade local válida.",
                                       parent=popup)
                return
            if not assunto:
                messagebox.showwarning("Atenção", "Informe o assunto do e-mail.",
                                       parent=popup)
                return

            destinatario = self._EMAILS_UNIDADES[municipio]
            popup.destroy()
            threading.Thread(
                target=self._enviar_email,
                args=(remetente, senha, destinatario, municipio, assunto, mensagem, ids),
                daemon=True
            ).start()

        ctk.CTkButton(
            btns, text="Enviar",
            height=46, corner_radius=12,
            fg_color=_VERDE, hover_color=_VERDE_H,
            font=("Segoe UI", 13, "bold"), text_color="#fff",
            command=_confirmar,
        ).grid(row=0, column=0, padx=(0, 6), sticky="ew")

        ctk.CTkButton(
            btns, text="Cancelar",
            height=46, corner_radius=12,
            fg_color=AppTheme.BG_INPUT,
            hover_color=AppTheme.BG_APP,
            font=("Segoe UI", 13),
            text_color=AppTheme.TXT_MAIN,
            command=popup.destroy,
        ).grid(row=0, column=1, padx=(6, 0), sticky="ew")

    def _enviar_email(self, remetente, senha, destinatario, municipio,
                      assunto, mensagem, ids):
        """Envia e-mail com os PDFs como anexo via Gmail SMTP (roda em thread)."""
        try:
            msg = MIMEMultipart()
            msg["From"]    = remetente
            msg["To"]      = destinatario
            msg["Subject"] = assunto
            msg.attach(MIMEText(mensagem, "plain", "utf-8"))

            anexos_ok = 0
            for analise_id in ids:
                resultado = self.controller.repo.buscar_pdf_binario(analise_id)
                if not resultado or not resultado.get("pdf_conteudo"):
                    continue
                nome = resultado.get("nome_pdf", f"processo_{analise_id}.pdf")
                parte = MIMEApplication(resultado["pdf_conteudo"], _subtype="pdf")
                parte.add_header("Content-Disposition", "attachment",
                                 filename=nome)
                msg.attach(parte)
                anexos_ok += 1

            if anexos_ok == 0:
                self.after(0, messagebox.showerror, "Erro",
                           "Nenhum PDF encontrado no banco para os registros selecionados.")
                return

            with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
                smtp.login(remetente, senha)
                smtp.sendmail(remetente, destinatario, msg.as_bytes())

            self.after(0, messagebox.showinfo, "Sucesso",
                       f"E-mail enviado com sucesso para {municipio}!\n"
                       f"Destinatário: {destinatario}\n"
                       f"Anexos: {anexos_ok} PDF(s)")

        except smtplib.SMTPAuthenticationError:
            self.after(0, messagebox.showerror, "Erro de Autenticação",
                       "E-mail ou senha incorretos.\n"
                       "Use uma Senha de App do Google (não a senha da conta).\n"
                       "Acesse: myaccount.google.com → Segurança → Senhas de app")
        except smtplib.SMTPException as e:
            self.after(0, messagebox.showerror, "Erro SMTP", str(e))
        except Exception as e:
            self.after(0, messagebox.showerror, "Erro",
                       f"Falha ao enviar e-mail:\n{str(e)}")

    def _enviar_email_form(self):
        """
        Abre popup de envio de e-mail para a unidade local.
        O município é detectado automaticamente pelo campo preenchido,
        mas o usuário pode trocar antes de enviar.
        """
        # Tenta detectar o município digitado no formulário
        muni_digitado = self.municipio.get().strip().upper()

        # Encontra a chave mais próxima no dicionário (busca parcial)
        muni_detectado = ""
        for chave in self._EMAILS_UNIDADES:
            if muni_digitado and muni_digitado in chave.upper():
                muni_detectado = chave
                break
            if chave.upper() == muni_digitado:
                muni_detectado = chave
                break

        popup = ctk.CTkToplevel(self)
        popup.title("Enviar E-mail para Unidade Local")
        popup.geometry("580x820")
        popup.resizable(False, False)
        popup.configure(fg_color=AppTheme.BG_APP)
        popup.grab_set()
        popup.after(0, lambda: popup.geometry(
            f"580x820+"
            f"{popup.winfo_screenwidth()  // 2 - 290}+"
            f"{popup.winfo_screenheight() // 2 - 410}"
        ))

        # ScrollableFrame para caber tudo sem cortar
        outer = ctk.CTkScrollableFrame(popup, fg_color="transparent",
                                       scrollbar_button_color=AppTheme.BG_INPUT)
        outer.pack(fill="both", expand=True)

        wrap = ctk.CTkFrame(outer, fg_color="transparent")
        wrap.pack(fill="both", expand=True, padx=28, pady=24)

        ctk.CTkLabel(wrap, text="Enviar E-mail para Unidade Local",
                     font=("Segoe UI", 18, "bold"),
                     text_color=AppTheme.TXT_MAIN).pack(anchor="w",
                                                         pady=(0, 16))

        def _entry_card(parent, placeholder, show=""):
            e = ctk.CTkEntry(
                parent, placeholder_text=placeholder,
                height=40, corner_radius=10,
                font=("Segoe UI", 12),
                fg_color=AppTheme.BG_INPUT,
                border_color=AppTheme.BG_INPUT,
                text_color=AppTheme.TXT_MAIN,
            )
            if show:
                e.configure(show=show)
            return e

        # ── Remetente ──────────────────────────────────────────────
        card_rem = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                corner_radius=14)
        card_rem.pack(fill="x", pady=(0, 12))

        ctk.CTkLabel(card_rem, text="Remetente",
                     font=("Segoe UI", 12, "bold"),
                     text_color=_MUTED).pack(anchor="w", padx=20,
                                              pady=(14, 6))

        ctk.CTkLabel(card_rem, text="E-mail:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_rem = _entry_card(card_rem, "seu.email@gmail.com")
        entry_rem.pack(fill="x", padx=20, pady=(2, 8))

        ctk.CTkLabel(card_rem, text="Senha de App (Gmail):",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_senha = _entry_card(card_rem, "Senha de app do Google",
                                  show="*")
        entry_senha.pack(fill="x", padx=20, pady=(2, 14))

        # ── Destinatário ───────────────────────────────────────────
        card_dest = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                 corner_radius=14)
        card_dest.pack(fill="x", pady=(0, 12))

        ctk.CTkLabel(card_dest, text="Unidade Local Destinatária",
                     font=("Segoe UI", 12, "bold"),
                     text_color=_MUTED).pack(anchor="w", padx=20,
                                              pady=(14, 6))

        municipios = sorted(self._EMAILS_UNIDADES.keys())
        combo_muni = ctk.CTkComboBox(
            card_dest,
            values=municipios,
            height=40, corner_radius=10,
            font=("Segoe UI", 12),
            fg_color=AppTheme.BG_INPUT,
            border_color=AppTheme.BG_INPUT,
            button_color=_VERDE,
            button_hover_color=_VERDE_H,
            text_color=AppTheme.TXT_MAIN,
            dropdown_fg_color=AppTheme.BG_CARD,
            dropdown_text_color=AppTheme.TXT_MAIN,
        )
        combo_muni.pack(fill="x", padx=20, pady=(0, 6))

        # Pré-seleciona o município detectado
        if muni_detectado:
            combo_muni.set(muni_detectado)

        lbl_email_dest = ctk.CTkLabel(
            card_dest, text="",
            font=("Segoe UI", 11), text_color=_VERDE)
        lbl_email_dest.pack(anchor="w", padx=20, pady=(0, 14))

        def _on_muni(choice):
            email = self._EMAILS_UNIDADES.get(choice, "")
            lbl_email_dest.configure(text=f"  ✉  {email}")

        combo_muni.configure(command=_on_muni)
        # Mostrar e-mail imediatamente se município foi detectado
        if muni_detectado:
            _on_muni(muni_detectado)

        # ── Assunto ────────────────────────────────────────────────
        card_msg = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                corner_radius=14)
        card_msg.pack(fill="x", pady=(0, 16))

        ctk.CTkLabel(card_msg, text="Assunto:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20, pady=(14, 2))
        entry_assunto = _entry_card(card_msg, "Ex: Devolução de processo")

        # Preenche assunto automaticamente com dados do formulário
        nome_form = self.nome.get().strip()
        memo_form = self.memorando.get().strip()
        tipo_reg  = self.radio.get()
        assunto_auto = (
            f"{'Devolução' if tipo_reg == 'dev' else 'Inscrição/Renovação'}"
            f"{' — ' + nome_form if nome_form else ''}"
            f"{' — Memo ' + memo_form if memo_form else ''}"
        )
        entry_assunto.insert(0, assunto_auto)
        entry_assunto.pack(fill="x", padx=20, pady=(0, 14))

        # ── Dados do Processo ──────────────────────────────────────
        card_proc = ctk.CTkFrame(wrap, fg_color=AppTheme.BG_CARD,
                                 corner_radius=14)
        card_proc.pack(fill="x", pady=(0, 16))

        ctk.CTkLabel(card_proc, text="Dados do Processo",
                     font=("Segoe UI", 12, "bold"),
                     text_color=_MUTED).pack(anchor="w", padx=20,
                                              pady=(14, 6))

        # Nome
        ctk.CTkLabel(card_proc, text="Nome:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_nome_proc = _entry_card(card_proc, "Nome do produtor")
        entry_nome_proc.insert(0, self.nome.get().strip())
        entry_nome_proc.pack(fill="x", padx=20, pady=(2, 8))

        # CPF
        ctk.CTkLabel(card_proc, text="CPF:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_cpf_proc = _entry_card(card_proc, "CPF")
        entry_cpf_proc.insert(0, self.cpf.get().strip())
        entry_cpf_proc.pack(fill="x", padx=20, pady=(2, 8))

        # Memorando
        ctk.CTkLabel(card_proc, text="Memorando:",
                     font=("Segoe UI", 11), text_color=_MUTED).pack(
            anchor="w", padx=20)
        entry_memo_proc = _entry_card(card_proc, "Memorando")
        entry_memo_proc.insert(0, self.memorando.get().strip())
        entry_memo_proc.pack(fill="x", padx=20, pady=(2, 8))

        # Categoria e Motivo — só no modo devolução
        entry_cat_proc  = None
        txt_mot_proc    = None

        if tipo_reg == "dev":
            ctk.CTkLabel(card_proc, text="Categoria:",
                         font=("Segoe UI", 11), text_color=_MUTED).pack(
                anchor="w", padx=20)
            entry_cat_proc = _entry_card(card_proc, "Categoria")
            entry_cat_proc.insert(0, self.categoria.get().strip())
            entry_cat_proc.pack(fill="x", padx=20, pady=(2, 8))

            ctk.CTkLabel(card_proc, text="Motivo da Devolução:",
                         font=("Segoe UI", 11), text_color=_MUTED).pack(
                anchor="w", padx=20)
            txt_mot_proc = ctk.CTkTextbox(
                card_proc, height=80, corner_radius=10,
                font=("Segoe UI", 12),
                fg_color=AppTheme.BG_INPUT,
                border_width=0,
                text_color=AppTheme.TXT_MAIN,
            )
            mot_atual = self.motivo.get("1.0", END).strip()
            if mot_atual:
                txt_mot_proc.insert("1.0", mot_atual)
            txt_mot_proc.pack(fill="x", padx=20, pady=(2, 14))
        else:
            ctk.CTkLabel(card_proc, text="Tipo:",
                         font=("Segoe UI", 11), text_color=_MUTED).pack(
                anchor="w", padx=20)
            entry_tipo_proc = _entry_card(card_proc, "Tipo")
            entry_tipo_proc.insert(0, self.tipo.get().strip())
            entry_tipo_proc.pack(fill="x", padx=20, pady=(2, 14))

        # ── Botões ─────────────────────────────────────────────────
        btns = ctk.CTkFrame(wrap, fg_color="transparent")
        btns.pack(fill="x")
        btns.grid_columnconfigure((0, 1), weight=1)

        def _confirmar():
            remetente = entry_rem.get().strip()
            senha     = entry_senha.get().strip()
            municipio = combo_muni.get().strip()
            assunto   = entry_assunto.get().strip()

            if not remetente or not senha:
                messagebox.showwarning("Atenção",
                    "Preencha o e-mail e a senha do remetente.",
                    parent=popup)
                return
            if not municipio or municipio not in self._EMAILS_UNIDADES:
                messagebox.showwarning("Atenção",
                    "Selecione uma unidade local válida.",
                    parent=popup)
                return
            if not assunto:
                messagebox.showwarning("Atenção",
                    "Informe o assunto do e-mail.",
                    parent=popup)
                return

            destinatario = self._EMAILS_UNIDADES[municipio]

            # Lê os campos editáveis do card de processo
            nome_f = entry_nome_proc.get().strip()
            cpf_f  = entry_cpf_proc.get().strip()
            memo_f = entry_memo_proc.get().strip()
            muni_f = self.municipio.get().strip()

            if tipo_reg == "dev":
                cat_f = entry_cat_proc.get().strip() if entry_cat_proc else ""
                mot_f = txt_mot_proc.get("1.0", END).strip() if txt_mot_proc else ""
                corpo = (
                    f"Prezados,\n\n"
                    f"Segue processo para devolução:\n\n"
                    f"Nome:       {nome_f}\n"
                    f"CPF:        {cpf_f}\n"
                    f"Município:  {muni_f}\n"
                    f"Memorando:  {memo_f}\n"
                    f"Categoria:  {cat_f}\n"
                    f"Motivo:     {mot_f}\n\n"
                    f"Atenciosamente,\n{self.usuario}"
                )
            else:
                tipo_f = self.tipo.get().strip()
                corpo = (
                    f"Prezados,\n\n"
                    f"Segue processo para {'inscrição' if tipo_f == 'INSC' else 'renovação'}:\n\n"
                    f"Nome:       {nome_f}\n"
                    f"CPF:        {cpf_f}\n"
                    f"Município:  {muni_f}\n"
                    f"Memorando:  {memo_f}\n"
                    f"Tipo:       {tipo_f}\n\n"
                    f"Atenciosamente,\n{self.usuario}"
                )

            popup.destroy()
            threading.Thread(
                target=self._enviar_email_simples,
                args=(remetente, senha, destinatario,
                      municipio, assunto, corpo),
                daemon=True,
            ).start()

        ctk.CTkButton(
            btns, text="Enviar",
            height=46, corner_radius=12,
            fg_color=_VERDE, hover_color=_VERDE_H,
            font=("Segoe UI", 13, "bold"), text_color="#fff",
            command=_confirmar,
        ).grid(row=0, column=0, padx=(0, 6), sticky="ew")

        ctk.CTkButton(
            btns, text="Cancelar",
            height=46, corner_radius=12,
            fg_color=AppTheme.BG_INPUT,
            hover_color=AppTheme.BG_APP,
            font=("Segoe UI", 13),
            text_color=AppTheme.TXT_MAIN,
            command=popup.destroy,
        ).grid(row=0, column=1, padx=(6, 0), sticky="ew")

    def abrir_historico(self):
        HistoricoView(self, self.usuario, self.controller).grab_set()
